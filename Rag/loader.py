"""
RAG模块 - 文档加载器
支持: txt, pdf, doc, docx
"""
import os
from pathlib import Path
from typing import List
from langchain_community.document_loaders import (
    TextLoader,
    PyPDFLoader,
)

try:
    from langchain_community.document_loaders import UnstructuredWordDocumentLoader
    HAS_UNSTRUCTURED = True
except ImportError:
    HAS_UNSTRUCTURED = False
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document


class DocxLoader:
    """基于 python-docx 的文档加载器"""
    
    def __init__(self, file_path: str):
        self.file_path = file_path
    
    def load(self) -> List[Document]:
        """加载 docx 文件"""
        import docx
        doc = docx.Document(self.file_path)
        
        # 提取所有段落文本
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        
        # 合并为完整文本
        full_text = '\n'.join(paragraphs)
        
        # 创建 Document 对象
        from langchain_core.documents import Document
        return [Document(
            page_content=full_text,
            metadata={"source": self.file_path}
        )]


class LegalDocLoader:
    """法律文档加载器"""
    
    SUPPORTED_EXTENSIONS = ['.txt', '.pdf', '.doc', '.docx']
    
    def __init__(self, docs_dir: str = None):
        self.docs_dir = Path(docs_dir) if docs_dir else None
        self.splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=100,
            separators=["\n\n", "\n", "。", "！", "？", " ", ""]
        )
    
    def load_file(self, file_path: str) -> List[Document]:
        """加载单个文件"""
        path = Path(file_path)
        ext = path.suffix.lower()
        
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"不支持的文件格式: {ext}")
        
        loader = self._get_loader(path, ext)
        docs = loader.load()
        
        # 分割文档
        if len(docs) > 1:
            return self.splitter.split_documents(docs)
        else:
            return self.splitter.split_documents(docs)
    
    def _get_loader(self, path: Path, ext: str):
        """根据扩展名获取加载器"""
        if ext == '.txt':
            return TextLoader(str(path), encoding='utf-8')
        elif ext == '.pdf':
            return PyPDFLoader(str(path))
        elif ext in ['.doc', '.docx']:
            # 优先使用 python-docx，更稳定
            return self._create_docx_loader(path)
        else:
            raise ValueError(f"不支持的格式: {ext}")
    
    def _create_docx_loader(self, path: Path):
        """创建 docx 加载器，优先使用 python-docx"""
        try:
            import docx
            # 使用自定义加载器
            return DocxLoader(str(path))
        except ImportError:
            if HAS_UNSTRUCTURED:
                return UnstructuredWordDocumentLoader(str(path))
            else:
                raise ImportError("请安装 python-docx: pip install python-docx")
    
    def load_directory(self, directory: str) -> List[Document]:
        """加载目录下所有支持的文件"""
        docs = []
        dir_path = Path(directory)
        
        for file_path in dir_path.rglob('*'):
            if file_path.is_file() and file_path.suffix.lower() in self.SUPPORTED_EXTENSIONS:
                try:
                    docs.extend(self.load_file(str(file_path)))
                except Exception:
                    pass
        
        return docs
    
    def get_supported_formats(self) -> List[str]:
        """获取支持的文件格式"""
        return self.SUPPORTED_EXTENSIONS